"""Document text extraction (TASK-017).

01_REQUIREMENTS.md § Evidence Document Ingestion, Explicitly Forbidden
Behavior: "The system must never execute, render, or open uploaded files in a
way that could trigger embedded active content (macros, scripts) — extraction
must use passive parsing libraries only."

Every library used here is a pure-Python parser that reads structure and never
evaluates it:

* `pypdf` — parses the PDF object graph. It does not run JavaScript actions,
  does not follow launch actions, and does not fetch remote resources.
* `python-docx` / `openpyxl` — read the OOXML parts as XML. Neither runs VBA;
  `openpyxl` is opened with `data_only=True` so cached values are read rather
  than formulas being evaluated.
* `pytesseract` — OCR on a rasterised image. It reads pixels.

Extraction never raises out of `extract`. Every failure becomes an
`ExtractionResult` with a message safe to show an auditor, because a crashed
worker leaves a row stuck in `processing` — the exact state 02_ARCHITECTURE.md
§7.5 requires the system not to have.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

# A page or sheet yielding less than this is treated as having no usable text.
_MIN_USEFUL_CHARS = 3


@dataclass
class ExtractedSection:
    """A unit of text with a human-readable location for citation."""

    location: str
    text: str


@dataclass
class ExtractionResult:
    success: bool
    sections: list[ExtractedSection] = field(default_factory=list)
    error: str | None = None

    @property
    def full_text(self) -> str:
        return "\n\n".join(f"[{s.location}]\n{s.text}" for s in self.sections)


def extract(content: bytes, mime_type: str) -> ExtractionResult:
    """Extract text from an uploaded document. Never raises."""
    try:
        if mime_type == "application/pdf":
            return _extract_pdf(content)
        if mime_type.endswith("wordprocessingml.document"):
            return _extract_docx(content)
        if mime_type.endswith("spreadsheetml.sheet"):
            return _extract_xlsx(content)
        if mime_type in ("image/png", "image/jpeg"):
            return _extract_image(content)
    except Exception:
        # The stack trace goes to the log; the auditor gets a plain sentence.
        # 02_ARCHITECTURE.md §7.7 — internal detail never reaches the client.
        logger.exception("Extraction failed for mime_type=%s", mime_type)
        return ExtractionResult(
            success=False,
            error="This document could not be read. Please review it manually.",
        )
    return ExtractionResult(success=False, error=f"No extractor for type {mime_type}.")


def _extract_pdf(content: bytes) -> ExtractionResult:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(io.BytesIO(content))
    except PdfReadError:
        return ExtractionResult(
            success=False, error="This PDF appears to be corrupt and could not be opened."
        )

    if reader.is_encrypted:
        # 01_REQUIREMENTS.md Edge Cases: rejected with a specific error asking
        # for an unprotected copy. The system does not attempt to guess or
        # brute-force the password, including trying the empty password —
        # attempting to defeat a protection measure is not this tool's business.
        return ExtractionResult(
            success=False,
            error=(
                "This PDF is password-protected. Please obtain an unprotected copy "
                "from the client and upload that instead."
            ),
        )

    sections: list[ExtractedSection] = []
    for number, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception:
            logger.warning("Could not extract page %d of a PDF", number)
            continue
        if len(text) >= _MIN_USEFUL_CHARS:
            sections.append(ExtractedSection(location=f"page {number}", text=text))

    if not sections:
        # A scanned PDF has pages but no text layer. OCR-ing it would mean
        # rasterising the PDF, which needs a renderer — a far larger dependency
        # than anything else here, and one that does interpret page content.
        # ponytail: image-only PDFs route to manual review instead. Add
        # pdf2image + poppler if scanned PDFs turn out to be common in practice;
        # the auditor can meanwhile upload a page image, which this module OCRs.
        logger.info("PDF has no text layer on any of %d pages", len(reader.pages))
        return ExtractionResult(
            success=False,
            error=(
                "No readable text was found in this PDF. It may be a scan. Please review "
                "it manually, or upload the relevant pages as images."
            ),
        )

    return ExtractionResult(success=True, sections=sections)


def _extract_docx(content: bytes) -> ExtractionResult:
    import docx

    document = docx.Document(io.BytesIO(content))

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    sections: list[ExtractedSection] = []
    if paragraphs:
        sections.append(ExtractedSection(location="document body", text="\n".join(paragraphs)))

    # Tables carry most of the structured evidence in a policy document, so they
    # are extracted separately rather than being flattened into the body.
    for index, table in enumerate(document.tables, start=1):
        rows = [
            " | ".join(cell.text.strip() for cell in row.cells)
            for row in table.rows
            if any(cell.text.strip() for cell in row.cells)
        ]
        if rows:
            sections.append(ExtractedSection(location=f"table {index}", text="\n".join(rows)))

    if not sections:
        return ExtractionResult(success=False, error="This document contains no readable text.")
    return ExtractionResult(success=True, sections=sections)


def _extract_xlsx(content: bytes) -> ExtractionResult:
    import openpyxl

    # data_only=True reads cached values instead of formulas: no expression is
    # evaluated, and the auditor sees what the client saw.
    workbook = openpyxl.load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    sections: list[ExtractedSection] = []
    try:
        for sheet in workbook.worksheets:
            rows: list[str] = []
            for row_number, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    rows.append(f"row {row_number}: " + " | ".join(cells))
            if rows:
                sections.append(
                    ExtractedSection(location=f"sheet '{sheet.title}'", text="\n".join(rows))
                )
    finally:
        workbook.close()

    if not sections:
        return ExtractionResult(success=False, error="This spreadsheet contains no data.")
    return ExtractionResult(success=True, sections=sections)


def _extract_image(content: bytes) -> ExtractionResult:
    """OCR an image. Screenshots are common evidence for configuration state."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return ExtractionResult(
            success=False,
            error="Image text extraction is not available on this server.",
        )

    try:
        image = Image.open(io.BytesIO(content))
        image.load()  # decode now, so a malformed image fails here and is caught
    except Exception:
        return ExtractionResult(
            success=False, error="This image could not be opened and may be corrupt."
        )

    try:
        text = pytesseract.image_to_string(image).strip()
    except pytesseract.TesseractNotFoundError:
        # A configuration problem, not a bad document. Saying so lets an
        # operator fix it instead of the auditor re-uploading forever.
        logger.error("Tesseract is not installed; image OCR is unavailable")
        return ExtractionResult(
            success=False,
            error="Image text extraction is not configured on this server.",
        )

    if len(text) < _MIN_USEFUL_CHARS:
        return ExtractionResult(
            success=False,
            error="No readable text was found in this image. Please review it manually.",
        )
    return ExtractionResult(success=True, sections=[ExtractedSection("image", text)])


def chunk_sections(
    sections: list[ExtractedSection], *, max_chars: int = 1500, overlap: int = 150
) -> list[tuple[str, str]]:
    """Split sections into embedding-sized chunks of (location, text).

    Chunks never span sections, so a chunk's `location` is always truthful —
    that is what lets a Finding cite "page 3" and be right. The overlap keeps a
    statement that straddles a boundary intact in at least one chunk.
    """
    chunks: list[tuple[str, str]] = []
    for section in sections:
        text = section.text.strip()
        if not text:
            continue
        if len(text) <= max_chars:
            chunks.append((section.location, text))
            continue

        start = 0
        part = 1
        while start < len(text):
            end = min(start + max_chars, len(text))
            if end < len(text):
                # Prefer to break at a paragraph or sentence boundary so a chunk
                # does not end mid-sentence, which reads badly in a citation.
                boundary = max(
                    text.rfind("\n", start + max_chars // 2, end),
                    text.rfind(". ", start + max_chars // 2, end),
                )
                if boundary > start:
                    end = boundary + 1
            chunks.append((f"{section.location} (part {part})", text[start:end].strip()))
            if end >= len(text):
                break
            start = max(start + 1, end - overlap)
            part += 1
    return [(location, text) for location, text in chunks if text]
