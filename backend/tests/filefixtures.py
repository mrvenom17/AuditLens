"""Builders for real files of each accepted and rejected type.

The upload and extraction tests need genuine bytes rather than stubs: the
security control under test is content-type inspection by magic bytes
(05_SECURITY.md §10.4), so a placeholder string would prove nothing about
whether a real DOCX is distinguished from a real ZIP.
"""

from __future__ import annotations

import io
import zipfile


def valid_pdf(text: str = "Firewall configuration export. Rule set reviewed.") -> bytes:
    """A small, well-formed PDF with a real text layer."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    pdf.drawString(72, 800, text)
    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def multipage_pdf(pages: list[str]) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    for text in pages:
        pdf.drawString(72, 800, text)
        pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def password_protected_pdf(password: str = "s3cret") -> bytes:
    """01_REQUIREMENTS.md Edge Cases: the system must reject these with a
    specific error rather than attempting to open them."""
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(io.BytesIO(valid_pdf("Protected content.")))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(password)

    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def corrupt_pdf() -> bytes:
    """A file that passes the magic-byte check but is not a parseable PDF.

    This is the case that must degrade to `extraction_failed` rather than crash
    the worker (TASK-017).
    """
    return b"%PDF-1.4\n" + b"\x00\xff" * 400 + b"\ntrailer garbage no xref"


def truncated_pdf() -> bytes:
    """A real PDF cut off mid-stream — a plausible result of a failed transfer."""
    return valid_pdf()[:120]


def valid_docx(paragraphs: list[str] | None = None) -> bytes:
    import docx

    document = docx.Document()
    for line in paragraphs or ["Access control policy v3.", "Reviewed annually."]:
        document.add_paragraph(line)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Control"
    table.cell(0, 1).text = "Owner"
    table.cell(1, 0).text = "MFA enforcement"
    table.cell(1, 1).text = "Platform team"

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def valid_xlsx(rows: list[list[str]] | None = None) -> bytes:
    import openpyxl

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Firewall"
    for row in rows or [["Rule", "Action"], ["inbound 443", "allow"], ["inbound 23", "deny"]]:
        sheet.append(row)

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def valid_png(text_like: bool = True) -> bytes:
    """A PNG. `text_like` renders large black-on-white text so OCR has
    something to find on a machine where tesseract is installed."""
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (600, 200), "white")
    if text_like:
        draw = ImageDraw.Draw(image)
        draw.text((20, 80), "MFA ENABLED FOR ALL ADMIN ACCOUNTS", fill="black")

    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def valid_jpeg() -> bytes:
    from PIL import Image

    image = Image.new("RGB", (120, 120), "white")
    buffer = io.BytesIO()
    image.save(buffer, format="JPEG")
    return buffer.getvalue()


# --- Files that must be rejected --------------------------------------------


def disguised_executable() -> bytes:
    """A Windows PE binary. Uploaded as `report.pdf` with an
    `application/pdf` content-type, this is the case 01_REQUIREMENTS.md names
    directly: "Given a .exe file renamed to .pdf, when uploaded, content-type
    inspection rejects it with 400"."""
    return b"MZ\x90\x00\x03\x00\x00\x00\x04\x00\x00\x00\xff\xff\x00\x00" + b"\x00" * 512


def elf_executable() -> bytes:
    """A Linux binary — the same attack on a different host platform."""
    return b"\x7fELF\x02\x01\x01\x00" + b"\x00" * 512


def shell_script() -> bytes:
    return b"#!/bin/sh\nrm -rf /\n"


def plain_zip() -> bytes:
    """A ZIP that is not an OOXML document.

    DOCX and XLSX share the ZIP magic bytes, so the allow-list has to look
    inside the archive. A plain ZIP must not slip through on the header alone.
    """
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("notes.txt", "not a document")
    return buffer.getvalue()


def zip_containing_executable() -> bytes:
    """A ZIP carrying a binary, named `.docx`. Rejected for the same reason as
    a plain ZIP: neither `word/document.xml` nor `xl/workbook.xml` is present."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("payload.exe", disguised_executable())
    return buffer.getvalue()


def svg_with_script() -> bytes:
    """An SVG carrying script. Not on the allow-list at all, and worth an
    explicit test because SVG renders as an image but executes as a document —
    the classic stored-XSS upload."""
    return b'<svg xmlns="http://www.w3.org/2000/svg"><script>alert(1)</script></svg>'


def html_file() -> bytes:
    return b"<html><body><script>fetch('//evil.example')</script></body></html>"


def oversized(size_mb: int = 26) -> bytes:
    """A valid-looking PDF over the 25MB cap (05_SECURITY.md §10.4)."""
    return b"%PDF-1.4\n" + b"A" * (size_mb * 1024 * 1024)


def empty_file() -> bytes:
    return b""
